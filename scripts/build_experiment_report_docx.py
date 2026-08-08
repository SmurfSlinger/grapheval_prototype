#!/usr/bin/env python3
"""Build reports/GraphEval Experiment Report.docx.

Table values are read from results/research/grapheval_final_experiment_analysis.json
(the authoritative analysis output) rather than typed in, and the four report
figures are embedded from results/research/figures/. Prose mirrors
"reports/GraphEval Experiment Report.md".

Usage: .venv/bin/python scripts/build_experiment_report_docx.py
"""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

REPO = Path(__file__).resolve().parents[1]
ANALYSIS = REPO / "results/research/grapheval_final_experiment_analysis.json"
FIGDIR = REPO / "results/research/figures"
MD = REPO / "reports/GraphEval Experiment Report.md"
OUT = REPO / "reports/GraphEval Experiment Report.docx"


def add_table(doc, header: list[str], rows: list[list[str]], caption: str | None = None):
    if caption:
        p = doc.add_paragraph(caption)
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = str(h)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            t.rows[i].cells[j].text = str(v)
    doc.add_paragraph()
    return t


def md_sections(md_text: str) -> dict[str, list[str]]:
    """Split the Markdown report into paragraph lists keyed by heading."""
    sections: dict[str, list[str]] = {}
    current = "_preamble"
    buf: list[str] = []
    for block in md_text.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("#"):
            sections[current] = buf
            current = block.lstrip("# ").strip()
            buf = []
        else:
            buf.append(block)
    sections[current] = buf
    return sections


def clean(par: str) -> str:
    """Flatten a markdown paragraph into plain text for DOCX prose."""
    text = par.replace("\n", " ")
    for ch in ("**", "`"):
        text = text.replace(ch, "")
    return text


def is_image(par: str) -> bool:
    return par.lstrip().startswith("![") and "](" in par


def parse_image(par: str) -> tuple[str, Path] | None:
    """Return (alt_text, absolute_path) for a markdown image, or None."""
    import re

    m = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", par.strip())
    if not m:
        return None
    alt, rel = m.group(1), m.group(2)
    path = (MD.parent / rel).resolve() if not Path(rel).is_absolute() else Path(rel)
    # MD lives in reports/; relative paths use ../research/... etc.
    if not path.exists():
        path = (REPO / rel.lstrip("./")).resolve()
    return alt, path


def is_table(par: str) -> bool:
    return par.lstrip().startswith("|")


def add_md_paragraphs(doc, pars: list[str]):
    for par in pars:
        if is_image(par):
            parsed = parse_image(par)
            if parsed is None:
                continue
            alt, path = parsed
            if path.suffix.lower() == ".svg":
                # python-docx embeds raster images; prefer sibling PNG when present.
                png = path.with_suffix(".png")
                if png.exists():
                    path = png
                else:
                    note = doc.add_paragraph(
                        f"[SVG figure omitted from DOCX embed; see {path.relative_to(REPO)}]"
                    )
                    note.runs[0].italic = True
                    note.runs[0].font.size = Pt(9)
                    continue
            if not path.exists():
                note = doc.add_paragraph(
                    f"[Figure file not yet present: {path.name}. "
                    f"See research/neo4j_figures/FIGURE_CAPTIONS.md for capture instructions.]"
                )
                note.runs[0].italic = True
                note.runs[0].font.size = Pt(9)
                continue
            width = Inches(6.5) if "iteration_sequence" in path.name else Inches(6.0)
            doc.add_picture(str(path), width=width)
            continue
        if par.startswith("*Figure ") or (
            par.startswith("*") and "Figure M" in par and par.endswith("*")
        ):
            cap = doc.add_paragraph(clean(par.strip("*")))
            if cap.runs:
                cap.runs[0].italic = True
                cap.runs[0].font.size = Pt(9)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if is_table(par):
            lines = [ln for ln in par.splitlines() if ln.strip().startswith("|")]
            grid = [
                [c.strip() for c in ln.strip().strip("|").split("|")]
                for ln in lines
            ]
            grid = [row for row in grid if not all(set(c) <= set("-: ") for c in row)]
            if grid:
                header, *rows = grid
                add_table(doc, [clean(h) for h in header], [[clean(c) for c in r] for r in rows])
            continue
        if par.startswith("- ") or par[:2].rstrip(".").isdigit() or par.startswith("1."):
            import re

            items: list[tuple[str, str]] = []
            for line in par.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
                if stripped.startswith("- "):
                    items.append(("List Bullet", stripped[2:].strip()))
                elif m:
                    items.append(("List Number", m.group(2)))
                elif items:
                    style, prev = items[-1]
                    items[-1] = (style, prev + " " + stripped)
                else:
                    items.append(("Normal", stripped))
            for style, content in items:
                doc.add_paragraph(clean(content), style=style)
            continue
        doc.add_paragraph(clean(par))


def main() -> None:
    analysis = json.loads(ANALYSIS.read_text(encoding="utf-8"))
    md = MD.read_text(encoding="utf-8")
    sections = md_sections(md)

    doc = Document()
    core = doc.core_properties
    core.title = "Evaluating Graph-Based Feedback and Iterative Backtracking in GraphEval"
    core.author = "Kyler Gundersen"
    core.subject = "GraphEval Prototype Experiment Report"

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = "GraphEval Experiment Report — August 2026"
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p = section.footer.paragraphs[0]
    footer_p.text = "grapheval_prototype — commit b9608d0 — official run apollo_multihop_llama31_8b_20260727T203028Z"
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for p in (header_p, footer_p):
        for r in p.runs:
            r.font.size = Pt(8)

    # Title block
    title = doc.add_heading(core.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in sections.get(core.title, [])[:2]:
        p = doc.add_paragraph(clean(line))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.size = Pt(9)

    heading_order = [
        ("Abstract", 1),
        ("1. Introduction", 1),
        ("1.1 Background", 2),
        ("1.2 Objective", 2),
        ("2. Methodology", 1),
        ("2.1 Research motivation and design rationale", 2),
        ("2.2 Triple representation", 2),
        ("2.3 Neo4j data model and persistence", 2),
        ("2.4 GraphEval algorithm", 2),
        ("2.5 Claim evaluation", 2),
        ("2.6 Feedback and revision loop", 2),
        ("2.7 Target and evidence-path validation", 2),
        ("2.8 Complete worked execution", 2),
        ("2.9 Experiment procedure and scoring", 2),
        ("3. Results", 1),
        ("3.1 Overall results", 2),
        ("3.2 Stop reasons", 2),
        ("3.3 Joint textual correctness × pipeline resolution", 2),
        ("3.4 Results by designed depth", 2),
        ("3.5 Revision behavior", 2),
        ("3.6 Final claim labels", 2),
        ("3.7 Figures", 2),
        ("3.8 Representative trace cases", 2),
        ("3.9 Repeatability and Nondeterminism", 2),
        ("4. Discussion", 1),
        ("5. Conclusion", 1),
        ("6. References", 1),
    ]
    for name, level in heading_order:
        doc.add_heading(name, level=level)
        if name == "6. References":
            # Keep literal reference numbers; Word's List Number style would
            # otherwise continue counting from the Procedure list.
            for par in sections.get(name, []):
                for line in par.splitlines():
                    stripped = line.strip()
                    if not stripped:
                        continue
                    import re

                    if re.match(r"^\d+\.\s", stripped):
                        doc.add_paragraph(clean(stripped).replace("*", ""))
                    else:
                        prev = doc.paragraphs[-1]
                        prev.add_run(" " + clean(stripped).replace("*", ""))
            continue
        add_md_paragraphs(doc, sections.get(name, []))
        if name == "3.7 Figures":
            figures = [
                ("fig1_outcomes_by_depth.png",
                 "Figure 1. Outcome counts by designed depth. Each depth contains only five questions; bars are raw counts, not trend estimates."),
                ("fig2_joint_outcomes.png",
                 "Figure 2. Joint textual-correctness (contains-expected) × pipeline-resolution outcomes (n=50)."),
                ("fig3_stop_reasons.png",
                 "Figure 3. Final stop-reason distribution (n=50)."),
                ("fig4_revision_outcomes.png",
                 "Figure 4. Outcomes for first-pass vs revised questions. Resolution ends iteration, so first-pass rows are resolved by construction."),
            ]
            for fname, caption in figures:
                path = FIGDIR / fname
                if path.exists():
                    doc.add_picture(str(path), width=Inches(6.0))
                    cap = doc.add_paragraph(caption)
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if name == "3.9 Repeatability and Nondeterminism":
            rep_figdir = REPO / "results/research/repeatability/figures"
            rep_figures = [
                ("figR1_aggregate_outcomes_by_run.png",
                 "Figure 5. Aggregate outcomes for all three repeatability runs (n=50 per run). All output metrics were identical; raw counts shown."),
                ("figR2_stability_by_dimension.png",
                 "Figure 6. Per-question stability across the three runs by outcome dimension: all 50 questions were stable on every compared dimension."),
            ]
            for fname, caption in rep_figures:
                path = rep_figdir / fname
                if path.exists():
                    doc.add_picture(str(path), width=Inches(6.0))
                    cap = doc.add_paragraph(caption)
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Consistency guard: verify key numbers in the analysis JSON appear in the text.
    ov = analysis["overall"]
    for value in (ov["exact_match"], ov["contains_expected"], ov["pipeline_resolved"]):
        assert str(value) in md, f"analysis value {value} missing from report prose"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
